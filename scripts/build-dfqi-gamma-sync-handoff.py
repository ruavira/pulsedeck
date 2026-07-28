#!/usr/bin/env python3
"""Build the facilitator-safe DFQI Gamma/PulseDeck pilot handoff.

The document intentionally contains no presenter key, remote URL, or local
credential-file content.
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "DFQI-Gamma-PulseDeck-Facilitator-Handoff.docx"

TEAL = "0F766E"
TEAL_PALE = "F0FDFA"
NAVY = "0B2545"
INK = "1F2937"
MUTED = "5B6472"
GRID = "D7E1E7"
LIGHT = "F2F4F7"
WHITE = "FFFFFF"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, *, name="Calibri", size=11, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_cell_border(cell, color=GRID, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def add_hyperlink(paragraph, text: str, url: str, color=TEAL):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    r_pr.append(run_color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_numbering_definition(doc):
    numbering = doc.part.numbering_part.element
    abstract_num_id = max(
        [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
        or [0]
    ) + 1
    num_id = max(
        [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
        or [0]
    ) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_num_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, indent, spacing])
    level.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbered_step(doc, num_id: int, title: str, detail: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    lead = p.add_run(f"{title}. ")
    set_run_font(lead, bold=True, color=NAVY)
    body = p.add_run(detail)
    set_run_font(body)


def add_status_callout(doc, label: str, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.line_spacing = 1.2
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), TEAL_PALE)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "20")
    left.set(qn("w:color"), TEAL)
    borders.append(left)
    p_pr.append(borders)
    lead = p.add_run(f"{label}: ")
    set_run_font(lead, bold=True, color=TEAL)
    body = p.add_run(text)
    set_run_font(body, color=NAVY)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

# compact_reference_guide token map with a named PulseDeck teal accent override.
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = rgb(INK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, before, after, color in (
    ("Heading 1", 16, 18, 10, TEAL),
    ("Heading 2", 13, 14, 7, TEAL),
    ("Heading 3", 12, 10, 5, NAVY),
):
    style = doc.styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = rgb(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
header.paragraph_format.space_after = Pt(0)
set_run_font(header.add_run("DFQI Part 1  |  Facilitator pilot handoff"), size=9, color=MUTED, bold=True)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer.paragraph_format.space_before = Pt(0)
set_run_font(footer.add_run("PulseDeck Gamma Sync  |  28 July 2026"), size=8.5, color=MUTED)

kicker = doc.add_paragraph()
kicker.paragraph_format.space_before = Pt(8)
kicker.paragraph_format.space_after = Pt(0)
set_run_font(kicker.add_run("LIVE TRAINING INTEGRATION"), size=10, color=TEAL, bold=True)

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(8)
set_run_font(title.add_run("Gamma + PulseDeck\nFacilitator Handoff"), size=28, color=NAVY, bold=True)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(18)
set_run_font(
    subtitle.add_run("DFQI Part 1, Sessions 1-3  |  One-control browser pilot"),
    size=13.5,
    color=MUTED,
)

add_status_callout(
    doc,
    "Pilot status",
    "PASSED. Gamma navigation now advances PulseDeck, opens the mapped interaction, displays live results inside Gamma, and removes the interaction panel on the next content card.",
)

doc.add_heading("What is ready", level=1)
for lead, detail in (
    ("Protected Gamma copies", "The three source presentations remain untouched; all edits and testing use named pilot copies."),
    ("Combined PulseDeck deck", "155 mapped Gamma cards, including 21 paced discussion interactions (7 per session)."),
    ("Chrome integration", "PulseDeck Gamma Sync version 0.2.0 is installed and enabled in the current Chrome profile."),
    ("Seamless live layer", "A right-side PulseDeck panel appears only on join and activity cards. Ordinary Gamma cards stay unobstructed."),
):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    set_run_font(p.add_run(f"{lead}: "), bold=True, color=NAVY)
    set_run_font(p.add_run(detail))

doc.add_heading("Protected Gamma pilot presentations", level=1)
pilot_links = (
    (
        "Session 1 - Signal or Noise?",
        "https://gamma.app/docs/DFQI-Part-1-Session-1-Signal-or-Noise-PulseDeck-Live-Pilot-cy7x7cox38l68ru",
    ),
    (
        "Session 2 - From Metric to Chart",
        "https://gamma.app/docs/DFQI-Part-1-Session-2-From-Metric-to-Chart-PulseDeck-Live-Pilot-dc0sfypwingbucz",
    ),
    (
        "Session 3 - From Chart to Boardroom",
        "https://gamma.app/docs/DFQI-Part-1-Session-3-From-Chart-to-Boardroom-PulseDeck-Live-Pilo-f7h6h98pib25r7f",
    ),
)
for label, url in pilot_links:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    add_hyperlink(p, label, url)

doc.add_heading("Before each live delivery", level=1)
num_id = add_numbering_definition(doc)
add_numbered_step(
    doc,
    num_id,
    "Start the combined PulseDeck deck",
    "Create a fresh live session. Keep its complete Remote URL private; it contains presenter access.",
)
add_numbered_step(
    doc,
    num_id,
    "Connect Gamma Sync",
    "Open the extension while a pilot Gamma is active, paste the complete Remote URL, and choose Connect live session. Confirm that it reports 155 Gamma cards mapped.",
)
add_numbered_step(
    doc,
    num_id,
    "Open the participant join",
    "Show the first Gamma card long enough for participants to scan the join QR or enter the session code.",
)
add_numbered_step(
    doc,
    num_id,
    "Present from Gamma",
    "Use Gamma's normal arrow-key navigation. Do not operate the PulseDeck remote unless the automatic sync fails.",
)
add_numbered_step(
    doc,
    num_id,
    "Run a 3-minute rehearsal",
    "Open one poll, submit at least one test response, advance one card, and confirm that the panel disappears before participants join.",
)

doc.add_heading("What the facilitator will see", level=1)
doc.add_heading("On an activity card", level=2)
p = doc.add_paragraph(
    "The extension badge reads ON. PulseDeck advances to the mapped activity, opens voting automatically, and overlays live results on the right side of Gamma. The Gamma prompt remains visible on the left."
)
p.paragraph_format.space_after = Pt(6)

doc.add_heading("On the next content card", level=2)
p = doc.add_paragraph(
    "One Gamma arrow key advances both systems. PulseDeck returns to the neutral show phase and the overlay is removed immediately."
)
p.paragraph_format.space_after = Pt(6)

doc.add_heading("Interaction pacing", level=1)
p = doc.add_paragraph(
    "Seven moments per session were selected to support discussion without turning the training into a game. The cadence is approximately one interaction every 8-9 minutes across the 3-hour sequence."
)
p.paragraph_format.space_after = Pt(6)

for session, mix in (
    ("Session 1 - Think", "3 polls, 3 open-text reflections, 1 confidence scale."),
    ("Session 2 - Chart", "4 polls, 1 open-text debrief, 1 word cloud, 1 confidence scale."),
    ("Session 3 - Communicate", "3 polls, 3 open-text decisions/reflections, 1 confidence scale."),
):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run(f"{session}: "), bold=True, color=NAVY)
    set_run_font(p.add_run(mix))

doc.add_page_break()
doc.add_heading("Pilot verification evidence", level=1)
table = doc.add_table(rows=1, cols=3)
set_table_geometry(table, [2480, 1120, 5760])
headers = ("Boundary", "Status", "Evidence")
for idx, value in enumerate(headers):
    cell = table.rows[0].cells[idx]
    set_cell_shading(cell, TEAL)
    set_cell_border(cell, color=TEAL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(value), size=10, color=WHITE, bold=True)

evidence_rows = (
    ("Mapping", "PASS", "155 unique Gamma cards; 21 activities; all three protected pilot slugs."),
    ("Gamma to API", "PASS", "Activity card krnae3syuwn1fpk advanced PulseDeck to slide 115 and phase open."),
    ("API to overlay", "PASS", "The mapped poll rendered inside Gamma presentation mode without a presenter key in the iframe."),
    ("Audience results", "PASS", "12 simulated participants produced 10 responses; counts and percentages updated live in Gamma."),
    ("Close/advance", "PASS", "One Right Arrow moved to card x97s8tut1ei5juq, set phase show, and removed the overlay."),
    ("Build quality", "PASS", "8 integration tests, ESLint, Next.js production build, and whitespace checks passed."),
)
for row_idx, values in enumerate(evidence_rows, start=1):
    row = table.add_row()
    cant_split = OxmlElement("w:cantSplit")
    row._tr.get_or_add_trPr().append(cant_split)
    cells = row.cells
    for idx, value in enumerate(values):
        set_cell_border(cells[idx])
        if row_idx % 2 == 0:
            set_cell_shading(cells[idx], LIGHT)
        p = cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(
            p.add_run(value),
            size=9.5,
            color=TEAL if idx == 1 else INK,
            bold=(idx == 1),
        )
set_table_geometry(table, [2480, 1120, 5760])

doc.add_heading("Fallback and recovery", level=1)
for lead, detail in (
    ("Badge shows -", "The current Gamma card is not mapped. Confirm you opened a protected pilot copy, not a source deck."),
    ("Badge shows !", "A network or presenter-auth error occurred. Open the extension and reconnect with the current session's Remote URL."),
    ("Overlay does not appear", "Refresh the Gamma tab once, then choose Sync this card now in the extension."),
    ("Automatic sync stops", "Use the existing PulseDeck phone remote as the manual fallback; participant responses remain in the same live session."),
    ("Chrome was restarted", "Reconnect the extension. The presenter key is intentionally stored only for the Chrome session and is cleared on exit."),
):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    set_run_font(p.add_run(f"{lead}: "), bold=True, color=NAVY)
    set_run_font(p.add_run(detail))

add_status_callout(
    doc,
    "Security note",
    "Never paste the PulseDeck Remote URL into Gamma, a participant message, a screenshot, or this handoff. Public overlays use the credential-free deck embed; presenter control remains inside the extension's temporary session storage.",
)

doc.core_properties.title = "DFQI Gamma + PulseDeck Facilitator Handoff"
doc.core_properties.subject = "One-control live browser pilot and facilitator operating guide"
doc.core_properties.author = "Ruavira Collective Inc."
doc.core_properties.keywords = "DFQI, Gamma, PulseDeck, facilitator, live interaction"

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
