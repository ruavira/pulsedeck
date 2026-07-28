#!/usr/bin/env python3
"""Build the account-independent DFQI Gamma/PulseDeck continuity handoff.

The output deliberately excludes presenter credentials, private Remote URLs,
environment values, and browser/session storage.
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "handoff-2026-07-28-dfqi-gamma-pulsedeck.docx"

TEAL = "0F766E"
TEAL_PALE = "F0FDFA"
NAVY = "0B2545"
INK = "1F2937"
MUTED = "5B6472"
GRID = "D7E1E7"
LIGHT = "F2F4F7"
AMBER = "8A5A00"
AMBER_PALE = "FFF7E6"
WHITE = "FFFFFF"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run(run, *, size=11, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def shade_paragraph(paragraph, fill: str, border_color: str | None = None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border_color:
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "20")
        left.set(qn("w:color"), border_color)
        borders.append(left)
        p_pr.append(borders)


def add_callout(doc, label: str, text: str, *, caution=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.2
    shade_paragraph(p, AMBER_PALE if caution else TEAL_PALE, AMBER if caution else TEAL)
    set_run(p.add_run(f"{label}: "), bold=True, color=AMBER if caution else TEAL)
    set_run(p.add_run(text), color=NAVY)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def border_cell(cell, color=GRID, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_table(doc, headers, rows, widths, *, status_col=None):
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, TEAL)
        border_cell(cell, TEAL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(value), size=9.5, color=WHITE, bold=True)
    for row_num, values in enumerate(rows, start=1):
        row = table.add_row()
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            border_cell(cell)
            if row_num % 2 == 0:
                shade_cell(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            is_status = status_col is not None and idx == status_col
            set_run(p.add_run(value), size=9.2, color=TEAL if is_status else INK, bold=is_status)
    set_table_geometry(table, widths)
    return table


def add_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_id = max(
        [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))] or [0]
    ) + 1
    num_id = max([int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))] or [0]) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    for tag, val in (("start", "1"), ("numFmt", "decimal"), ("lvlText", "%1."), ("lvlJc", "left")):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:val"), val)
        level.append(node)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, indent, spacing])
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_step(doc, num_id, title, detail):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_node = OxmlElement("w:numId")
    num_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_node])
    set_run(p.add_run(f"{title}. "), bold=True, color=NAVY)
    set_run(p.add_run(detail))


def add_labeled(doc, label, detail, *, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    set_run(p.add_run(f"{label}: "), bold=True, color=color)
    set_run(p.add_run(detail))


def add_field(paragraph, instruction):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = OxmlElement("w:r")
    run.extend([begin, instr, separate, value, end])
    paragraph._p.append(run)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

# compact_reference_guide with a named PulseDeck teal accent override.
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = rgb(INK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, before, after, color in (
    ("Heading 1", 16, 18, 10, TEAL),
    ("Heading 2", 13, 14, 7, TEAL),
    ("Heading 3", 12, 10, 5, NAVY),
):
    style = doc.styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = rgb(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.paragraph_format.space_after = Pt(0)
set_run(header.add_run("DFQI continuity handoff  |  Gamma + PulseDeck"), size=9, color=MUTED, bold=True)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(footer.add_run("28 July 2026  |  Page "), size=8.5, color=MUTED)
add_field(footer, "PAGE")

# customer_pack opening pattern: left-aligned operational title + compact metadata.
kicker = doc.add_paragraph()
kicker.paragraph_format.space_before = Pt(8)
kicker.paragraph_format.space_after = Pt(0)
set_run(kicker.add_run("ACCOUNT-INDEPENDENT PROJECT CHECKPOINT"), size=10, color=TEAL, bold=True)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(6)
set_run(title.add_run("DFQI Gamma + PulseDeck\nContinuity Handoff"), size=27, color=NAVY, bold=True)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(16)
set_run(subtitle.add_run("For continuation in another Codex or Claude account"), size=13, color=MUTED)

metadata = add_table(
    doc,
    ("Item", "Current value"),
    (
        ("Repository", "github.com/ruavira/pulsedeck"),
        ("Merged implementation", "PR #14 · merge ce59d1b · implementation efa0468"),
        ("Production deck", "155 slides · 21 response interactions · 1 Join surface"),
        ("Current blocker", "Waiting for the final combined Gamma URL and stable 155-card sequence"),
    ),
    [2700, 6660],
)

add_callout(
    doc,
    "Status",
    "The browser integration is complete and merged. Do not create the final Thursday session until the combined Gamma deck has been remapped and validated.",
)

doc.add_heading("Progress log", level=1)
doc.add_heading("Completed", level=2)
for label, detail in (
    ("Browser integration", "Gamma navigation advances PulseDeck and injects/removes the public interaction overlay."),
    ("Production content", "The combined PulseDeck deck contains 155 mapped slides and seven response activities per session."),
    ("Pilot", "A live test passed with 12 simulated participants, 10 responses, live percentages, and overlay removal on the next content card."),
    ("Security", "Presenter credentials stay in gitignored local files and Chrome session storage; no private URL is stored in repository content."),
    ("Documentation", "Technical documentation, checkpoint notes, mapping data, and the facilitator guide are already in the repository."),
):
    add_labeled(doc, label, detail)

doc.add_heading("Remaining, in priority order", level=2)
num_id = add_numbering(doc)
for title_text, detail in (
    ("Receive the final combined Gamma URL", "The user is currently merging and finalizing the three presentations."),
    ("Inventory all final cards", "Record the combined document slug, each card ID, heading, and order. Expect 155 cards unless changes were intentional."),
    ("Update the mapping", "Replace the three pilot slugs and any changed card IDs in content/dfqi-gamma-sync.json."),
    ("Update production PulseDeck", "Run the update script using the existing private local credential file without printing or committing it."),
    ("Verify", "Run the Gamma-sync test suite, lint, production build, API pilot, and live Chrome browser pilot."),
    ("Rehearse Wednesday", "Create a fresh rehearsal session and verify one poll plus the next-card overlay removal."),
    ("Run Thursday clean", "After edits are frozen, create a new production session and reconnect the extension with its new private Remote URL."),
):
    add_step(doc, num_id, title_text, detail)

doc.add_page_break()
doc.add_heading("Repository orientation", level=1)
add_table(
    doc,
    ("Path", "Why it matters"),
    (
        ("DFQI-CONTINUITY.md", "Canonical account-independent state, next steps, security rules, and paste-ready resume prompt."),
        ("content/dfqi-gamma-sync.json", "Source of truth for three Gamma documents, 155 cards, exact card IDs, and all activity definitions."),
        ("integrations/gamma-sync-extension/", "Manifest V3 extension that watches Gamma, advances PulseDeck, and controls the overlay."),
        ("scripts/update-dfqi-gamma-sync-deck.mjs", "Updates the existing production PulseDeck deck after mapping changes."),
        ("scripts/pilot-dfqi-gamma-sync.mjs", "Creates and validates a test session against the PulseDeck API."),
        ("docs/gamma-sync.md", "Technical and operational integration documentation."),
        ("docs/checkpoints/gamma-sync-implementation.md", "Implementation checkpoint and browser-pilot evidence."),
    ),
    [3700, 5660],
)

doc.add_heading("Interaction inventory", level=1)
add_labeled(doc, "Total", "21 response-producing interactions plus the initial Join surface.")
add_labeled(doc, "Types", "10 polls, 7 open-text prompts, 3 confidence scales, and 1 word cloud.")
add_labeled(doc, "Session 1", "Gamma cards 2, 8, 14, 21, 31, 43 and 49; PulseDeck slides 2, 8, 14, 21, 31, 43 and 49.")
add_labeled(doc, "Session 2", "Gamma cards 6, 14, 21, 26, 36, 42 and 56; PulseDeck slides 57, 65, 72, 77, 87, 93 and 107.")
add_labeled(doc, "Session 3", "Gamma cards 9, 15, 18, 21, 28, 32 and 42; PulseDeck slides 116, 122, 125, 128, 135, 139 and 149.")

add_callout(
    doc,
    "Mapping warning",
    "Copying cards into a new Gamma changes the document slug and may create new card IDs. Never update only the URL and assume the old IDs remain correct.",
    caution=True,
)

doc.add_heading("Validation commands", level=1)
for command in ("npm run test:gamma-sync", "npm run lint", "npm run build", "git diff --check"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(3)
    shade_paragraph(p, LIGHT)
    set_run(p.add_run(command), size=10, color=NAVY, bold=True)

doc.add_heading("Open blockers", level=1)
add_table(
    doc,
    ("Blocker", "Owner", "Unblocks when"),
    (
        ("Final combined Gamma URL and stable sequence", "User", "All three sessions are merged and final edits are complete."),
        ("New document slug and final card IDs", "Next agent", "The final combined URL can be inspected."),
        ("Wednesday rehearsal session", "Next agent / user", "Production deck is updated with final mappings."),
        ("Thursday production session", "Facilitator", "Thursday content is frozen; create a clean session immediately before delivery."),
    ),
    [4000, 1600, 3760],
)

doc.add_heading("Security rules", level=1)
for label, detail in (
    ("Never commit", ".env files, .pulsedeck-local, presenter keys, private Remote URLs, browser storage, API tokens, or raw credential output."),
    ("Never share", "The complete Remote URL with participants or in Gamma, screenshots, documents, issues, or pull requests."),
    ("Do not stage", "The local docs/guide-assets directory; it contains temporary real-site captures and remains intentionally untracked."),
    ("Always scan", "The staged file list and staged patch for credentials before every commit and push."),
):
    add_labeled(doc, label, detail, color=AMBER if label != "Always scan" else TEAL)

doc.add_page_break()
doc.add_heading("Paste-ready resume prompt", level=1)
prompt = (
    "Continue the DFQI Gamma + PulseDeck integration from DFQI-CONTINUITY.md in the repository root. "
    "Read that file first, followed by docs/gamma-sync.md, docs/checkpoints/gamma-sync-implementation.md, "
    "content/dfqi-gamma-sync.json, and integrations/gamma-sync-extension/. PR #14 is already merged into main.\n\n"
    "The user has finalized one combined Gamma deck. Its URL is: [PASTE FINAL GAMMA URL HERE]. "
    "Inspect that deck, enumerate all cards and final card IDs, reconcile it against the existing 155-card "
    "mapping, update the mapping and production PulseDeck deck, run tests/lint/build, and execute a live "
    "browser pilot. Preserve the 21 interaction definitions unless the user requests a change. Never expose "
    "or commit presenter credentials, .env files, .pulsedeck-local contents, or a private Remote URL. Use a "
    "separate rehearsal session for Wednesday and a fresh production session for Thursday."
)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.18)
p.paragraph_format.right_indent = Inches(0.08)
p.paragraph_format.space_after = Pt(14)
p.paragraph_format.line_spacing = 1.2
shade_paragraph(p, TEAL_PALE, TEAL)
set_run(p.add_run(prompt), size=10.5, color=NAVY)

doc.add_heading("Memory checkpoint", level=1)
add_labeled(doc, "Repository memory", "The same current-state summary was saved to ~/.Codex/projects/-Users-ayodejisamuels/memory/dfqi-gamma-pulsedeck.md on the original workstation.")
add_labeled(doc, "Local-only video", "/Users/ayodejisamuels/Downloads/PulseDeck/DFQI-Gamma-PulseDeck-Video-Guide.mp4 is available on the original workstation but is not required for continuation.")
add_labeled(doc, "Credential availability", "The original workstation retains gitignored PulseDeck credential files. A fresh clone elsewhere requires presenter access to be supplied securely by the user.")

add_callout(
    doc,
    "Definition of done",
    "The combined Gamma URL is mapped, all 155 cards are verified, production PulseDeck is updated, automated checks pass, the live Chrome pilot passes, Wednesday uses a rehearsal session, and Thursday uses a new clean session.",
)

doc.core_properties.title = "DFQI Gamma + PulseDeck Continuity Handoff"
doc.core_properties.subject = "Account-independent project checkpoint and resume instructions"
doc.core_properties.author = "Ruavira Collective Inc."
doc.core_properties.keywords = "DFQI, Gamma, PulseDeck, continuity, handoff, browser integration"

doc.save(OUTPUT)
print(OUTPUT)
